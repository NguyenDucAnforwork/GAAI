"""
Full 20-question GAIA Level 1 test runner (concurrent).

Runs the first 20 Level 1 questions from test_data/metadata.jsonl through the
workflow concurrently (Gemini 2.5 Flash for agents, gpt-4o-mini formatter).
Budget cap: $0.10 (only the formatter hits paid OpenAI; everything else is the
free Gemini tier, so real cost is a fraction of a cent).
"""
import asyncio
import json
import os
import re
import sys
import time
import datetime

# Disable Langfuse network overhead during batch eval. Skip Gemini key
# validation — only the video question uses Gemini, and it tries keys itself,
# so we don't want validation pings burning the scarce free-tier quota.
os.environ.setdefault("LANGFUSE_DISABLED", "true")
os.environ.setdefault("GEMINI_VALIDATE", "false")

from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv(), override=True)

sys.path.insert(0, os.path.dirname(__file__))

CONCURRENCY   = 6      # OpenAI primary has no per-minute cap at this volume
PER_Q_TIMEOUT = 200.0  # seconds per question (DeepSeek R1 reasoning is slow)
COST_CEILING  = 0.10

# ── helpers ────────────────────────────────────────────────────────────────────
def load_level1(path: str, n: int = 20):
    data = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            if line.strip():
                data.append(json.loads(line))
    return [x for x in data if x.get("Level") == 1][:n]

# GAIA validation files (public mirror, since the scoring API serves no files
# in this environment and the official dataset is gated).
import requests, base64
GAIA_FILE_BASE = "https://huggingface.co/datasets/asteriadyt/2023/resolve/main/validation"

def fetch_gaia_file(file_name: str) -> bytes:
    r = requests.get(f"{GAIA_FILE_BASE}/{file_name}",
                     headers={"User-Agent": "Mozilla/5.0"}, timeout=40)
    r.raise_for_status()
    return r.content

def extract_file_text(file_name: str, content: bytes) -> str:
    """Return text extracted from a docx/xlsx attachment (empty if unsupported)."""
    import io
    if file_name.lower().endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if file_name.lower().endswith((".xlsx", ".xls")):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    out.append(" | ".join(cells))
        return "\n".join(out)
    return ""

async def solve_image_question(question: str, file_name: str, content: bytes) -> str:
    """Answer an image-based question with a vision model (gpt-4o)."""
    from langchain_openai import ChatOpenAI
    from langchain_core.messages import HumanMessage
    b64 = base64.b64encode(content).decode()
    mime = "image/png" if file_name.lower().endswith(".png") else "image/jpeg"
    vis = ChatOpenAI(model="gpt-4o", api_key=os.getenv("OPENAI_API_KEY"),
                     temperature=0.0, max_tokens=1800)
    msg = HumanMessage(content=[
        {"type": "text", "text": question + "\n\nAnalyze the image carefully and "
         "end with a single line 'FINAL ANSWER: <answer>'."},
        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
    ])
    return (await vis.ainvoke([msg])).content

def extract_final_answer(response: str) -> str:
    m = re.search(r"FINAL ANSWER:\s*(.*?)(?:\n|$)", response, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    lines = [l.strip() for l in response.strip().split("\n") if l.strip()]
    return lines[-1] if lines else ""

def evaluate(predicted: str, expected: str) -> bool:
    """Honest GAIA-style match. Numbers must be exactly equal (not substring,
    so 17000 != 17); logic symbols (¬ → ↔ ∨ ∧) are preserved so different
    formulas don't collapse to equal."""
    def norm(t):
        # strip only basic punctuation/quotes; KEEP logic & math symbols
        t = re.sub(r"[\"'.,;:!?]", "", t.lower())
        return re.sub(r"\s+", " ", t).strip()
    p, e = norm(predicted), norm(expected)
    if not p:
        return False
    if p == e:
        return True
    # Numeric expected → require exact numeric equality with some number in p.
    e_nums = re.findall(r"-?\d+\.?\d*", e)
    if e_nums and re.fullmatch(r"-?\d+\.?\d*", e.replace(",", "")):
        try:
            target = float(e_nums[0])
            for x in re.findall(r"-?\d+\.?\d*", p.replace(",", "")):
                if abs(float(x) - target) < 1e-9:
                    return True
        except ValueError:
            pass
        return False
    # Text expected → accept if the full expected phrase appears in prediction.
    if e in p:
        return True
    return False

# ── main ───────────────────────────────────────────────────────────────────────
async def main():
    from langgraph_ver.workflow import OptimizedGAAIWorkflow

    questions = load_level1("test_data/metadata.jsonl", n=20)
    print(f"Loaded {len(questions)} Level 1 questions | concurrency={CONCURRENCY}\n")

    workflow = OptimizedGAAIWorkflow()
    sem = asyncio.Semaphore(CONCURRENCY)

    async def run_one(idx, q):
        async with sem:
            task_id  = q["task_id"]
            question = q["Question"]
            expected = str(q["Final answer"])
            file_name = (q.get("file_name") or "").strip()
            t0 = time.time()
            try:
                # Resolve any attachment from the public GAIA mirror.
                img_content = None
                if file_name:
                    try:
                        content = await asyncio.to_thread(fetch_gaia_file, file_name)
                        if file_name.lower().endswith((".png", ".jpg", ".jpeg")):
                            img_content = content
                        else:
                            text = extract_file_text(file_name, content)
                            if text:
                                question = question + "\n\n[Attached document contents]\n" + text
                    except Exception as fe:
                        print(f"file fetch failed for {file_name}: {fe}")

                if img_content is not None:
                    # Image questions go to a vision model directly.
                    response = await asyncio.wait_for(
                        solve_image_question(question, file_name, img_content),
                        timeout=PER_Q_TIMEOUT)
                else:
                    response = await asyncio.wait_for(
                        workflow.process_query(question), timeout=PER_Q_TIMEOUT)
            except asyncio.TimeoutError:
                response = "FINAL ANSWER: TIMEOUT"
            except Exception as e:
                response = f"FINAL ANSWER: ERROR ({e})"
            elapsed = time.time() - t0
            predicted = extract_final_answer(response)
            ok = evaluate(predicted, expected)
            mark = "✓" if ok else "✗"
            print(f"  {mark} Q{idx+1:02d} [{task_id[:8]}] ({elapsed:5.1f}s) "
                  f"exp={expected[:30]!r} got={predicted[:40]!r}")
            return {
                "task_id": task_id, "question": question, "expected": expected,
                "predicted": predicted, "response": response,
                "elapsed_time": round(elapsed, 1), "is_correct": ok,
            }

    t_start = time.time()
    results = await asyncio.gather(*[run_one(i, q) for i, q in enumerate(questions)])
    wall = time.time() - t_start

    correct = sum(1 for r in results if r["is_correct"])
    total = len(results)
    accuracy = correct / total if total else 0

    print(f"\n{'='*60}")
    print(f"RESULTS: {correct}/{total} correct = {accuracy*100:.1f}%")
    print(f"Wall-clock: {wall:.1f}s ({wall/60:.1f} min)")
    print(f"{'='*60}\n")
    print("Breakdown:")
    for i, r in enumerate(results):
        s = "✓" if r["is_correct"] else "✗"
        print(f"  {s} Q{i+1:02d} [{r['task_id'][:8]}] exp={r['expected'][:28]!r:30s} got={r['predicted'][:40]!r}")

    os.makedirs("results", exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out = f"results/full_test_{ts}.json"
    with open(out, "w", encoding="utf-8") as f:
        json.dump({"accuracy": accuracy, "correct": correct, "total": total,
                   "wall_seconds": round(wall, 1), "results": results}, f, indent=2)
    print(f"\nResults saved → {out}")

if __name__ == "__main__":
    asyncio.run(main())
